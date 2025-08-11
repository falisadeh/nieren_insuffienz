#%%
"""
04_time_to_AKI_ehrapy.py
=========================
Fragestellung: "Wie effektiv unterstützt das Framework ehrapy die Identifikation von Risikofaktoren für AKI
bei Kindern nach Herzoperationen anhand eines angereicherten Routinedatensatzes?"

Ansatz (ehrapy-first):
- **ehrapy/AnnData** ist die **zentrale Datenstruktur**. Alle abgeleiteten Variablen und Metadaten
  werden in `adata.obs` bzw. `adata.uns` abgelegt (Provenance!).
- Visualisierungen/Modelle verwenden stabile Ökosystem-Pakete (`lifelines`, `matplotlib`, `statsmodels/sklearn`).
- Ergebnis: Reproduzierbare Pipeline, die methodisch in ehrapy eingebettet bleibt.

Dieser Schritt (S1): Zeit-zu-Ereignis (Kaplan–Meier) für AKI ≤ 7 Tage
- Definiert **Index-OP** je AKI (letzte OP vor AKI im 0–7-Tage-Fenster), um Doppelzählungen zu vermeiden.
- Erzeugt `obs`-Variablen: `time_0_7`, `event_idx`.
- Schreibt Parameter/Provenance nach `adata.uns['survival_0_7']`.
- Exportiert CSV + 2 Plots; speichert Version der AnnData-Datei (H5AD) als S1.

Voraussetzungen
---------------
- Conda-Env: ehrapy_env (Python ≥ 3.10)
- Pakete: ehrapy, anndata, pandas, numpy, lifelines, matplotlib

Installation (falls nötig):
  conda install -c conda-forge ehrapy anndata pandas numpy lifelines matplotlib

"""
from __future__ import annotations

import os
from dataclasses import dataclass
from typing import Dict, Tuple

import numpy as np
import pandas as pd
import matplotlib.pyplot as plt

# ehrapy/AnnData laden (robust mit Fallback)
try:
    import ehrapy as ep  # type: ignore
except Exception:
    ep = None  # Fallback erlaubt, Kern bleibt AnnData

from anndata import read_h5ad, AnnData

try:
    from lifelines import KaplanMeierFitter
except Exception as e:
    raise SystemExit(
        "lifelines fehlt. Bitte installieren:\n"
        "  conda install -c conda-forge lifelines\n\n"
        f"Technischer Hinweis: {e}"
    )


# ---------------------------------
# 1) KONFIGURATION
# ---------------------------------
@dataclass
class Config:
    PATH_H5AD: str = (
        "/Users/fa/Library/Mobile Documents/com~apple~CloudDocs/cs-transfer/aki_ops_master.h5ad"
    )
    SAVE_DIR: str = (
        "/Users/fa/Library/Mobile Documents/com~apple~CloudDocs/cs-transfer/Diagramme"
    )
    OUT_H5AD: str = (
        "/Users/fa/Library/Mobile Documents/com~apple~CloudDocs/cs-transfer/aki_ops_master_S1_survival.h5ad"
    )
    WINDOW_DAYS: int = 7

CFG = Config()


# ---------------------------------
# 2) HILFSFUNKTIONEN (ehrapy-first)
# ---------------------------------

def load_adata() -> AnnData:
    """Lädt die AnnData-Struktur. Bevorzugt `ep.io.read_h5ad`, sonst Fallback `anndata.read_h5ad`.

    *Warum so?* Damit bleibt die Pipeline **ehrapy-first**, ist aber robust, falls einzelne I/O-Funktionen
    in Ihrer ehrapy-Version fehlen.
    """
    if ep is not None and hasattr(ep, "io") and hasattr(ep.io, "read_h5ad"):
        adata = ep.io.read_h5ad(CFG.PATH_H5AD)  # type: ignore[attr-defined]
    else:
        adata = read_h5ad(CFG.PATH_H5AD)
    return adata


def _ensure_datetime(df: pd.DataFrame, cols: Tuple[str, ...]) -> pd.DataFrame:
    """Robuste Datums-Konvertierung mit `errors='coerce'` (fehlerhafte Einträge -> NaT)."""
    df = df.copy()
    for c in cols:
        if c in df.columns:
            df[c] = pd.to_datetime(df[c], errors="coerce")
    return df


def annotate_provenance(adata: AnnData, key: str, meta: Dict) -> None:
    """Speichert Parameter/Provenance in `adata.uns[key]`.

    Beispiele: Fenstergröße, Versionshinweise, Anzahl Ereignisse.
    """
    if adata.uns is None:
        adata.uns = {}
    adata.uns[key] = meta


def define_event_index_on_adata(adata: AnnData) -> None:
    """Definiert den **Index-OP** je AKI-Episode und erzeugt `event_idx` in `obs`.

    Logik:
    - Kandidaten = OPs mit `AKI_linked_0_7 == 1` und gültigem `AKI_Start`.
    - Auswahl = letzte OP mit `Surgery_End` ≤ `AKI_Start` innerhalb der Gruppe (`PMID`, `AKI_Start`).
    - Ergebnis: `obs['event_idx']` ∈ {0,1}.
    """
    df = adata.obs.copy()

    # duration_hours ggf. aus X holen (falls nur 1 Var in var vorhanden)
    if "duration_hours" not in df.columns and "duration_hours" in adata.var_names:
        df = pd.concat([df, adata.to_df()[["duration_hours"]]], axis=1)

    # Typen/Zeiten vereinheitlichen
    df = _ensure_datetime(df, ("Surgery_Start", "Surgery_End", "AKI_Start"))
    if "days_to_AKI" in df.columns:
        df["days_to_AKI"] = pd.to_numeric(df["days_to_AKI"], errors="coerce")

    # Binärfelder harmonisieren
    for b in ("AKI_linked", "AKI_linked_0_7"):
        if b in df.columns:
            df[b] = (
                df[b].astype(str).str.strip().replace({"True": 1, "False": 0, "nan": np.nan, "None": np.nan})
            )
            df[b] = pd.to_numeric(df[b], errors="coerce").fillna(0).astype(int)

    # Offensichtliche Datenfehler entfernen (negatives Intervall nicht zulassen)
    if "days_to_AKI" in df.columns:
        df = df[df["days_to_AKI"].isna() | (df["days_to_AKI"] >= 0)].copy()

    # Ereignis-Index initialisieren
    df["event_idx"] = 0

    if "AKI_linked_0_7" in df.columns and df["AKI_linked_0_7"].sum() > 0:
        cand = df[(df["AKI_linked_0_7"] == 1) & df["AKI_Start"].notna()].copy()
        cand = cand[cand["Surgery_End"].notna()]
        cand = cand[cand["Surgery_End"] <= cand["AKI_Start"]]

        if not cand.empty:
            idx_rows = (
                cand.sort_values(["PMID", "AKI_Start", "Surgery_End"]).groupby(["PMID", "AKI_Start"], observed=True, as_index=False).tail(1).index
            )
            df.loc[idx_rows, "event_idx"] = 1

            multi_counts = cand.groupby(["PMID", "AKI_Start"], observed=True).size().rename("n_ops_in_window").reset_index()
            n_multi = int((multi_counts["n_ops_in_window"] > 1).sum())
        else:
            n_multi = 0
    else:
        n_multi = 0

    # Ergebnisse zurück nach adata.obs (Index beibehalten)
    adata.obs.loc[df.index, "event_idx"] = df["event_idx"].astype(int)

    # Provenance
    annotate_provenance(
        adata,
        key="survival_0_7_index",
        meta={
            "window_days": CFG.WINDOW_DAYS,
            "note": "Index-OP = letzte OP vor AKI im 0–7-Tage-Fenster",
            "n_events_marked": int(df["event_idx"].sum()),
            "n_total": int(df.shape[0]),
            "n_multi_ops_per_aki": int(n_multi),
        },
    )


def build_survival_on_adata(adata: AnnData) -> pd.DataFrame:
    """Erzeugt `time_0_7` (Tage) und nutzt `event_idx` aus `obs`.

    - Für Index-OPs: `time_0_7 = min(days_to_AKI, 7)`.
    - Für andere gelinkte OPs: `time_0_7 = min(days_to_AKI, 7)`, aber `event_idx = 0`.
    - Für unverbundene OPs: `time_0_7 = 7`, `event_idx = 0`.

    Speichert `time_0_7` dauerhaft in `adata.obs` und gibt den kompakten **Survival-DataFrame** zurück.
    """
    df = adata.obs.copy()

    if "days_to_AKI" not in df.columns:
        raise KeyError("'days_to_AKI' fehlt in adata.obs – bitte beim Preprocessing erzeugen.")

    if "event_idx" not in df.columns:
        raise KeyError("'event_idx' fehlt. Bitte zuerst define_event_index_on_adata() ausführen.")

    # Basis: zensiert bei WINDOW_DAYS
    df["time_0_7"] = float(CFG.WINDOW_DAYS)

    mask_evt = df["event_idx"] == 1
    df.loc[mask_evt, "time_0_7"] = np.clip(df.loc[mask_evt, "days_to_AKI"], 0, CFG.WINDOW_DAYS)

    mask_non_index_linked = (df.get("AKI_linked_0_7", 0) == 1) & (~mask_evt)
    df.loc[mask_non_index_linked, "time_0_7"] = np.clip(df.loc[mask_non_index_linked, "days_to_AKI"], 0, CFG.WINDOW_DAYS)

    # negative Zeiten vermeiden
    df.loc[df["time_0_7"] < 0, "time_0_7"] = 0.0

    # zurückschreiben
    adata.obs.loc[df.index, "time_0_7"] = df["time_0_7"].astype(float)

    # kompakten Survival-DF bilden
    keep = [c for c in ["time_0_7", "event_idx", "days_to_AKI", "AKI_linked_0_7", "Surgery_End", "AKI_Start", "PMID", "SMID", "Procedure_ID", "duration_hours", "Sex_norm"] if c in df.columns]
    surv = df[keep].copy()

    # Provenance
    annotate_provenance(
        adata,
        key="survival_0_7",
        meta={
            "window_days": CFG.WINDOW_DAYS,
            "n_rows": int(surv.shape[0]),
            "n_events": int((surv["event_idx"] == 1).sum()),
            "time_summary": {
                "min": float(surv["time_0_7"].min()),
                "median": float(surv["time_0_7"].median()),
                "max": float(surv["time_0_7"].max()),
            },
        },
    )

    return surv


def plot_km_and_cuminc(surv: pd.DataFrame, save_dir: str) -> Tuple[str, str]:
    """Speichert zwei Plots: KM (Überleben ohne AKI) und kumulative Inzidenz (1−S(t))."""
    os.makedirs(save_dir, exist_ok=True)

    # KM
    kmf = KaplanMeierFitter()
    T = surv["time_0_7"].astype(float)
    E = surv["event_idx"].astype(int)
    kmf.fit(T, event_observed=E, label="AKI ≤ 7 Tage")

    plt.figure(figsize=(7, 5))
    ax = kmf.plot(ci_show=True)
    ax.set_title("Kaplan–Meier: Zeit bis AKI ≤ 7 Tage ab OP-Ende")
    ax.set_xlabel("Tage seit OP-Ende")
    ax.set_ylabel("Überlebenswahrscheinlichkeit ohne AKI")
    ax.grid(True, alpha=0.3)
    km_path = os.path.join(save_dir, "KM_0_7_overall.png")
    plt.tight_layout(); plt.savefig(km_path, dpi=150); plt.close()

    # kumulative Inzidenz
    cuminc = 1.0 - kmf.survival_function_
    plt.figure(figsize=(7, 5))
    plt.step(cuminc.index.values, cuminc.values.flatten(), where="post")
    plt.title("Kumulative Inzidenz: AKI ≤ 7 Tage ab OP-Ende")
    plt.xlabel("Tage seit OP-Ende")
    plt.ylabel("Anteil mit AKI (1 − S(t))")
    plt.grid(True, alpha=0.3)
    ci_path = os.path.join(save_dir, "KM_0_7_cuminc.png")
    plt.tight_layout(); plt.savefig(ci_path, dpi=150); plt.close()

    return km_path, ci_path


# ---------------------------------
# 3) SUBGRUPPEN & BASISBESCHREIBUNG (S2)
# ---------------------------------

def compute_reop_and_duration_features(adata: AnnData) -> None:
    """Berechnet zusätzliche Variablen in `adata.obs`:
    - `is_reop` : 1 = Re-Operation (Patient hatte bereits eine frühere OP), sonst 0.
    - `duration_tertile` : kategorial (1/2/3) basierend auf `duration_hours`.

    Begründung:
    - Re-OPs sind klinisch relevant (höheres Risiko?).
    - Tertile der OP-Dauer erlauben eine robuste, nicht-parametrische Subgruppenbildung.
    """
    df = adata.obs.copy()

    # Sicherstellen, dass Zeitvariable für Sortierung existiert
    if "Surgery_End" in df.columns:
        df = df.copy()
        df["Surgery_End"] = pd.to_datetime(df["Surgery_End"], errors="coerce")
        # Reihenfolge je Patient
        df = df.sort_values(["PMID", "Surgery_End"]).copy()
        # Laufender Zähler pro Patient, beginnend bei 1
        df["op_seq"] = df.groupby("PMID", observed=True).cumcount() + 1
        df["is_reop"] = (df["op_seq"] > 1).astype(int)
        # zurückschreiben
        adata.obs.loc[df.index, "is_reop"] = df["is_reop"].astype(int)
    else:
        adata.obs["is_reop"] = 0  # konservativer Fallback

    # Dauer (in Stunden) prüfen
    if "duration_hours" not in adata.obs.columns and "duration_hours" in adata.var_names:
        # falls noch nicht in obs
        adata.obs = pd.concat([adata.obs, adata.to_df()[["duration_hours"]]], axis=1)

    if "duration_hours" in adata.obs.columns:
        # Tertile nur über valide Werte bilden
        dh = adata.obs["duration_hours"].astype(float)
        valid_mask = dh.notna()
        try:
            q = dh[valid_mask].quantile([1/3, 2/3]).values
            # 1: kurz, 2: mittel, 3: lang
            tertile = pd.Series(np.nan, index=adata.obs.index, dtype=float)
            tertile.loc[valid_mask & (dh <= q[0])] = 1
            tertile.loc[valid_mask & (dh > q[0]) & (dh <= q[1])] = 2
            tertile.loc[valid_mask & (dh > q[1])] = 3
            adata.obs["duration_tertile"] = tertile.astype("Int64")
        except Exception:
            adata.obs["duration_tertile"] = pd.Series(np.nan, index=adata.obs.index, dtype="Int64")
    else:
        adata.obs["duration_tertile"] = pd.Series(np.nan, index=adata.obs.index, dtype="Int64")


def km_by_group(surv: pd.DataFrame, group_col: str, save_dir: str, fname: str, order: list | None = None, label_map: dict | None = None) -> str:
    """Zeichnet KM-Kurven **stratifiziert** nach `group_col`.

    - `order` legt die Reihenfolge der Gruppen fest (optional).
    - `label_map` erlaubt schönere Legenden-Texte.
    """
    from lifelines import KaplanMeierFitter

    os.makedirs(save_dir, exist_ok=True)

    kmf = KaplanMeierFitter()

    plt.figure(figsize=(7, 5))

    groups = surv[group_col].dropna().unique().tolist()
    if order is not None:
        groups = [g for g in order if g in groups]

    for g in groups:
        sub = surv[surv[group_col] == g]
        if sub.empty:
            continue
        T = sub["time_0_7"].astype(float)
        E = sub["event_idx"].astype(int)
        label = label_map.get(g, str(g)) if label_map else str(g)
        kmf.fit(T, event_observed=E, label=label)
        kmf.plot(ci_show=False)

    plt.title(f"Kaplan–Meier (0–7 Tage) nach {group_col}")
    plt.xlabel("Tage seit OP-Ende")
    plt.ylabel("Überlebenswahrscheinlichkeit ohne AKI")
    plt.grid(True, alpha=0.3)
    plt.legend(title=group_col)

    out_path = os.path.join(save_dir, fname)
    plt.tight_layout(); plt.savefig(out_path, dpi=150); plt.close()

    return out_path


def make_table1_op_level(surv: pd.DataFrame, save_dir: str) -> str:
    """Erstellt eine einfache **Table 1** auf OP-Level, stratifiziert nach `event_idx` (0/1).

    Hinweis: `event_idx==1` bezeichnet **Index-OPs mit AKI**; `0` sind alle übrigen OPs (inkl. zensiert).
    Für die Bachelorarbeit ist das als *OP-basierte* Baseline sinnvoll, die *Patienten-basierte* Variante
    (erste OP pro Patient) können wir als Sensitivitätsanalyse ergänzen.
    """
    os.makedirs(save_dir, exist_ok=True)

    def _summary_num(x: pd.Series) -> dict:
        x = pd.to_numeric(x, errors="coerce")
        return {
            "n": int(x.notna().sum()),
            "mean": float(x.mean()),
            "sd": float(x.std()),
            "median": float(x.median()),
            "q1": float(x.quantile(0.25)),
            "q3": float(x.quantile(0.75)),
        }

    rows = []

    for grp, gdf in surv.groupby("event_idx", observed=True):
        label = "AKI-Index-OP" if grp == 1 else "Keine AKI-Index-OP"

        # Dauer
        if "duration_hours" in gdf.columns:
            rows.append({"Gruppe": label, "Variable": "duration_hours", **_summary_num(gdf["duration_hours"])})

        # Kategoriale: Geschlecht
        if "Sex_norm" in gdf.columns:
            vc = gdf["Sex_norm"].value_counts(dropna=False)
            n = int(len(gdf))
            for k, v in vc.items():
                rows.append({"Gruppe": label, "Variable": f"Sex_norm={k}", "n": int(v), "pct": float(100.0*v/n)})

        # Kategoriale: Re-OP
        if "is_reop" in gdf.columns:
            vc = gdf["is_reop"].value_counts(dropna=False)
            n = int(len(gdf))
            for k, v in vc.items():
                rows.append({"Gruppe": label, "Variable": f"is_reop={k}", "n": int(v), "pct": float(100.0*v/n)})

        # Kategoriale: Dauer-Tertile
        if "duration_tertile" in gdf.columns:
            vc = gdf["duration_tertile"].value_counts(dropna=False)
            n = int(len(gdf))
            for k, v in vc.items():
                rows.append({"Gruppe": label, "Variable": f"duration_tertile={k}", "n": int(v), "pct": float(100.0*v/n)})

    out = pd.DataFrame(rows)
    csv_path = os.path.join(save_dir, "table1_op_level.csv")
    out.to_csv(csv_path, index=False)
    return csv_path


# ---------------------------------
# 3) HAUPTABLAUF (S1)
# ---------------------------------

def main():
    # A) Laden (ehrapy-first, Fallback vorhanden)
    adata = load_adata()

    # B) Index-OPs definieren
    define_event_index_on_adata(adata)

    # C) Survival-Variablen aufbauen und zurück in adata.obs schreiben
    surv = build_survival_on_adata(adata)

    # D) CSV-Export (für externe Re-Use und QA)
    os.makedirs(CFG.SAVE_DIR, exist_ok=True)
    csv_path = os.path.join(CFG.SAVE_DIR, "survival_dataset_0_7.csv")
    surv.to_csv(csv_path, index=False)
    print(f"Survival-CSV gespeichert: {csv_path}")

    # E) Plots
    km_path, ci_path = plot_km_and_cuminc(surv, CFG.SAVE_DIR)
    print(f"Plots gespeichert: {km_path} | {ci_path}")

    # F) Versionierte H5AD-Datei als S1 persistieren (ehrapy/AnnData)
    adata.write_h5ad(CFG.OUT_H5AD)
    print(f"AnnData (S1) gespeichert: {CFG.OUT_H5AD}")

    print("\nFERTIG (S1): ehrapy-first Survival-Setup abgeschlossen.\n")


def main_s2():
    """S2: Subgruppenplots + Table 1 (ehrapy-integriert).

    Lädt, falls vorhanden, die S1-Datei (`CFG.OUT_H5AD`), ergänzt Subgruppen-Features und erzeugt
    stratifizierte KM-Kurven sowie eine einfache Table 1.
    """
    # Laden: bevorzugt die versionierte S1-Datei
    adata = read_h5ad(CFG.OUT_H5AD) if os.path.exists(CFG.OUT_H5AD) else load_adata()

    # Subgruppen-Features berechnen
    compute_reop_and_duration_features(adata)

    # Survival-DF aus obs zusammenstellen
    needed = ["time_0_7", "event_idx", "duration_hours", "Sex_norm", "is_reop", "duration_tertile"]
    surv_cols = [c for c in needed if c in adata.obs.columns]
    surv = adata.obs[surv_cols].copy()

    # KM nach Geschlecht
    if "Sex_norm" in surv.columns:
        _ = km_by_group(
            surv=surv,
            group_col="Sex_norm",
            save_dir=CFG.SAVE_DIR,
            fname="KM_0_7_by_sex.png",
            order=None,
            label_map=None,
        )

    # KM Erst-OP vs Re-OP
    if "is_reop" in surv.columns:
        _ = km_by_group(
            surv=surv,
            group_col="is_reop",
            save_dir=CFG.SAVE_DIR,
            fname="KM_0_7_by_reop.png",
            order=[0, 1],
            label_map={0: "Erst-OP", 1: "Re-OP"},
        )

    # KM nach Dauer-Tertilen
    if "duration_tertile" in surv.columns:
        _ = km_by_group(
            surv=surv,
            group_col="duration_tertile",
            save_dir=CFG.SAVE_DIR,
            fname="KM_0_7_by_duration_tertile.png",
            order=[1, 2, 3],
            label_map={1: "kurz", 2: "mittel", 3: "lang"},
        )

    # Table 1
    t1_path = make_table1_op_level(surv, CFG.SAVE_DIR)
    print(f"Table 1 gespeichert: {t1_path}")

    # Speichern (gleicher OUT-Pfad, aktualisierte obs)
    adata.write_h5ad(CFG.OUT_H5AD)
    print(f"AnnData (S1+S2) gespeichert: {CFG.OUT_H5AD}")


if __name__ == "__main__":
    main()
    main_s2()

# %%
# ===== S3: Table 1 mit Tests & Effektgrößen (ehrapy-first) =====
import math, json
try:
    from scipy import stats
    SCIPY_OK = True
except Exception:
    SCIPY_OK = False

def _num_summary(x: pd.Series):
    x = pd.to_numeric(x, errors="coerce").dropna()
    if x.empty:
        return {"n":0,"mean":np.nan,"sd":np.nan,"median":np.nan,"q1":np.nan,"q3":np.nan}
    return {
        "n": int(x.shape[0]),
        "mean": float(x.mean()),
        "sd": float(x.std(ddof=1)),
        "median": float(x.median()),
        "q1": float(x.quantile(0.25)),
        "q3": float(x.quantile(0.75)),
    }

def _r(x, nd=2):
    if x is None or (isinstance(x, float) and (np.isnan(x) or np.isinf(x))):
        return ""
    return f"{x:.{nd}f}"

def _binary_effects(ctab_2x2):
    # Reihenfolge: Zeilen=Kategorie (z.B. 0/1), Spalten=Gruppe (G1/G2)
    a,b = int(ctab_2x2.iloc[0,0]), int(ctab_2x2.iloc[0,1])
    c,d = int(ctab_2x2.iloc[1,0]), int(ctab_2x2.iloc[1,1])
    # OR (mit 0.5-Korrektur bei 0)
    aa,bb,cc,dd = a or 0.5, b or 0.5, c or 0.5, d or 0.5
    OR = (aa*dd)/(bb*cc)
    se = math.sqrt(1/aa + 1/bb + 1/cc + 1/dd)
    ci_lo, ci_hi = math.exp(math.log(OR)-1.96*se), math.exp(math.log(OR)+1.96*se)
    # RD für „erste Zeile“-Kategorie
    p1 = a/(a+c) if (a+c)>0 else np.nan
    p2 = b/(b+d) if (b+d)>0 else np.nan
    RD = p1 - p2 if not (np.isnan(p1) or np.isnan(p2)) else np.nan
    # p-Wert
    if SCIPY_OK:
        try:
            _, p = stats.fisher_exact(ctab_2x2.values)
            test = "Fisher"
        except Exception:
            chi2, p, _, _ = stats.chi2_contingency(ctab_2x2.values, correction=False)
            test = "Chi²"
    else:
        p, test = None, "—"
    return OR, (ci_lo, ci_hi), RD, p, test

def _kgt2_effect(ctab):
    # Kx2: Cramér’s V; Chi²-p
    if SCIPY_OK:
        chi2, p, _, _ = stats.chi2_contingency(ctab.values, correction=False)
        n = ctab.values.sum()
        k = min(ctab.shape)-1
        V = math.sqrt(chi2/(n*k)) if n>0 and k>0 else np.nan
        return V, p, "Chi²"
    else:
        return np.nan, None, "—"

def build_table1_with_stats(adata: AnnData, surv: pd.DataFrame, save_dir: str) -> pd.DataFrame:
    # Gruppenlabels
    gmap = {0: "Keine AKI-Index-OP", 1: "AKI-Index-OP"}
    surv = surv.copy()
    surv["Gruppe"] = surv["event_idx"].map(gmap)

    rows = []

    # --- Dauer (Stunden): Mediane/IQR + Mann-Whitney p (roh), Hedges g optional ---
    for g in [gmap[1], gmap[0]]:
        sx = surv.loc[surv["Gruppe"]==g, "duration_hours"].astype(float)
        s = _num_summary(sx)
        rows.append({"Variable":"Dauer (Stunden) – summary", "Gruppe":g, **s})
    # p-Werte (Mann-Whitney + Welch als Zusatz)
    x1 = surv.loc[surv["event_idx"]==1, "duration_hours"].astype(float).dropna()
    x0 = surv.loc[surv["event_idx"]==0, "duration_hours"].astype(float).dropna()
    if SCIPY_OK and x1.size>1 and x0.size>1:
        U, p_mw = stats.mannwhitneyu(x1, x0, alternative="two-sided")
        # Welch
        t, p_welch = stats.ttest_ind(x1, x0, equal_var=False)
    else:
        p_mw, p_welch = None, None
    # Hedges g
    if x1.size>1 and x0.size>1:
        m1, m0 = x1.mean(), x0.mean()
        s1, s0 = x1.std(ddof=1), x0.std(ddof=1)
        n1, n0 = x1.size, x0.size
        s_pooled = math.sqrt(((n1-1)*s1**2 + (n0-1)*s0**2) / (n1+n0-2)) if (n1+n0-2)>0 else np.nan
        if s_pooled and not np.isnan(s_pooled):
            d = (m1-m0)/s_pooled
            J = 1 - (3/(4*(n1+n0)-9)) if (n1+n0)>3 else 1.0
            g = J*d
        else:
            g = np.nan
    else:
        g = np.nan
    # Dauer – Ergebniszeile (druckfertig)
    dur_row = {
        "Variable": "Dauer (Stunden)",
        gmap[1]: f"{_r(x1.median())} [{_r(x1.quantile(.25))}; {_r(x1.quantile(.75))}] (n={x1.size})",
        gmap[0]: f"{_r(x0.median())} [{_r(x0.quantile(.25))}; {_r(x0.quantile(.75))}] (n={x0.size})",
        "Hedges g": _r(g,2),
        "p (Mann-Whitney)": _r(p_mw,3) if p_mw is not None else "",
        "p (Welch)": _r(p_welch,3) if p_welch is not None else "",
        "Hinweis": "Median[IQR] berichtet; nichtparametrischer Test (MW); Welch als Zusatz"
    }

    # --- Kategoriale: Geschlecht, Re-OP, Dauer-Tertile ---
    def cat_block(col, label):
        out = []
        ser = surv[col]
        if ser.isna().all():
            return out
        ctab = pd.crosstab(ser, surv["Gruppe"])
        # Binär?
        if ctab.shape[0]==2:
            OR, ci, RD, p, test = _binary_effects(ctab)
            for cat, row in ctab.iterrows():
                n1, n0 = int(row.get(gmap[1],0)), int(row.get(gmap[0],0))
                tot1 = int(ctab[gmap[1]].sum()); tot0 = int(ctab[gmap[0]].sum())
                out.append({
                    "Variable": f"{label}: {cat}",
                    gmap[1]: f"{n1} ({_r(100*n1/tot1,1)}%)" if tot1 else f"{n1}",
                    gmap[0]: f"{n0} ({_r(100*n0/tot0,1)}%)" if tot0 else f"{n0}",
                    "Effekt": f"OR {_r(OR,2)} [{_r(ci[0],2)}; {_r(ci[1],2)}]",
                    "p-Wert": _r(p,3) if p is not None else "",
                    "Test": test
                })
            out.append({
                "Variable": f"{label}: Risikodifferenz (G1−G2, erste Kategorie)",
                gmap[1]: f"n={int(ctab[gmap[1]].sum())}",
                gmap[0]: f"n={int(ctab[gmap[0]].sum())}",
                "Effekt": f"RD {_r(RD,3)}",
                "p-Wert": _r(p,3) if p is not None else "",
                "Test": test
            })
        else:
            V, p, test = _kgt2_effect(ctab)
            for cat, row in ctab.iterrows():
                n1, n0 = int(row.get(gmap[1],0)), int(row.get(gmap[0],0))
                tot1 = int(ctab[gmap[1]].sum()); tot0 = int(ctab[gmap[0]].sum())
                out.append({
                    "Variable": f"{label}: {cat}",
                    gmap[1]: f"{n1} ({_r(100*n1/tot1,1)}%)" if tot1 else f"{n1}",
                    gmap[0]: f"{n0} ({_r(100*n0/tot0,1)}%)" if tot0 else f"{n0}",
                    "Effekt": f"Cramérs V {_r(V,2)}",
                    "p-Wert": _r(p,3) if p is not None else "",
                    "Test": test
                })
        return out

    rows2 = []
    rows2 += cat_block("Sex_norm", "Geschlecht")
    if "is_reop" in surv.columns:
        rows2 += cat_block("is_reop", "Re-OP (ja=1, nein=0)")
    if "duration_tertile" in surv.columns:
        rows2 += cat_block("duration_tertile", "OP-Dauer (Tertile)")

    # Finale, druckfertige Tabelle
    final = [dur_row] + rows2
    table = pd.DataFrame(final)

    # Speichern
    os.makedirs(save_dir, exist_ok=True)
    out_csv = os.path.join(save_dir, "Table1_OP_level_with_stats.csv")
    table.to_csv(out_csv, index=False)
    print(f"Table 1 (mit Stats) gespeichert: {out_csv}")

    # In adata.uns['table1']
    adata.uns["table1"] = {
        "groups": [gmap[1], gmap[0]],
        "notes": {
            "duration_tests": "Mann-Whitney (primär), Welch t (Zusatz)",
            "categorical_tests": "Fisher/Chi²; Effekt: OR+CI (2x2) bzw. Cramérs V (>2 Kat.), RD zusätzlich"
        },
        "table": table.to_dict(orient="records")
    }
    return table

# %%
from anndata import read_h5ad
print("\n--- H5AD-Check ---")
adata_chk = read_h5ad(CFG.OUT_H5AD)
print(adata_chk)
print("uns keys:", list(adata_chk.uns.keys()))
print("table1 rows:", len(adata_chk.uns.get("table1", {}).get("table", [])))
#%%
# 05_table1_stats_ehrapy.py
# Erzeugt eine druckfertige Table 1 (CSV + optional DOCX) aus der H5AD
# und schreibt alles nach adata.uns['table1'] (ehrapy-first).

import os, math
import numpy as np
import pandas as pd
from anndata import read_h5ad

# ----- Pfade anpassen falls nötig -----
H5 = "/Users/fa/Library/Mobile Documents/com~apple~CloudDocs/cs-transfer/aki_ops_master_S1_survival.h5ad"
OUT_DIR = "/Users/fa/Library/Mobile Documents/com~apple~CloudDocs/cs-transfer/Diagramme"
CSV_OUT = os.path.join(OUT_DIR, "Table1_OP_level_with_stats.csv")
DOCX_OUT = os.path.join(OUT_DIR, "Table1_OP_level_with_stats.docx")

# SciPy optional (für p-Werte). Wenn nicht da, lassen wir p leer.
try:
    from scipy import stats
    SCIPY = True
except Exception:
    SCIPY = False

def r(x, nd=2):
    if x is None or (isinstance(x, float) and (np.isnan(x) or np.isinf(x))):
        return ""
    return f"{x:.{nd}f}"

def num_summary(x: pd.Series):
    x = pd.to_numeric(x, errors="coerce").dropna()
    if x.empty:
        return {"n":0,"mean":np.nan,"sd":np.nan,"median":np.nan,"q1":np.nan,"q3":np.nan}
    return {
        "n": int(x.shape[0]),
        "mean": float(x.mean()),
        "sd": float(x.std(ddof=1)),
        "median": float(x.median()),
        "q1": float(x.quantile(0.25)),
        "q3": float(x.quantile(0.75)),
    }

def binary_effects(ctab_2x2: pd.DataFrame):
    # Zeilen: Kategorie (z. B. 0/1), Spalten: Gruppen (G1/G0)
    a,b = int(ctab_2x2.iloc[0,0]), int(ctab_2x2.iloc[0,1])
    c,d = int(ctab_2x2.iloc[1,0]), int(ctab_2x2.iloc[1,1])
    # OR mit 0.5-Korrektur
    aa,bb,cc,dd = a or 0.5, b or 0.5, c or 0.5, d or 0.5
    OR = (aa*dd)/(bb*cc)
    se = math.sqrt(1/aa + 1/bb + 1/cc + 1/dd)
    ci_lo, ci_hi = math.exp(math.log(OR)-1.96*se), math.exp(math.log(OR)+1.96*se)
    # Risikodifferenz (erste Zeile als „positiv“ interpretiert)
    p1 = a/(a+c) if (a+c)>0 else np.nan
    p0 = b/(b+d) if (b+d)>0 else np.nan
    RD = p1 - p0 if not (np.isnan(p1) or np.isnan(p0)) else np.nan
    # p-Wert
    if SCIPY:
        try:
            _, p = stats.fisher_exact(ctab_2x2.values)
            test = "Fisher"
        except Exception:
            chi2, p, _, _ = stats.chi2_contingency(ctab_2x2.values, correction=False)
            test = "Chi²"
    else:
        p, test = None, "—"
    return OR, (ci_lo, ci_hi), RD, p, test

def kgt2_effect(ctab: pd.DataFrame):
    # Kx2: Cramér’s V; Chi²-p
    if SCIPY:
        chi2, p, _, _ = stats.chi2_contingency(ctab.values, correction=False)
        n = ctab.values.sum()
        k = min(ctab.shape) - 1
        V = math.sqrt(chi2/(n*k)) if n>0 and k>0 else np.nan
        return V, p, "Chi²"
    return np.nan, None, "—"

def build_table1(adata) -> pd.DataFrame:
    # Wir brauchen diese Spalten aus obs:
    need = ["event_idx","duration_hours","Sex_norm","is_reop","duration_tertile"]
    for c in need:
        if c not in adata.obs.columns:
            adata.obs[c] = np.nan

    surv = adata.obs[need].copy()
    gmap = {1: "AKI-Index-OP", 0: "Keine AKI-Index-OP"}
    surv["Gruppe"] = surv["event_idx"].map(gmap)

    rows = []

    # ---- Dauer (Stunden): MWU & Welch + Hedges g ----
    x1 = pd.to_numeric(surv.loc[surv["event_idx"]==1, "duration_hours"], errors="coerce").dropna()
    x0 = pd.to_numeric(surv.loc[surv["event_idx"]==0, "duration_hours"], errors="coerce").dropna()
    if SCIPY and x1.size>1 and x0.size>1:
        U, p_mw = stats.mannwhitneyu(x1, x0, alternative="two-sided")
        t, p_welch = stats.ttest_ind(x1, x0, equal_var=False)
    else:
        p_mw, p_welch = None, None

    # Hedges g
    g = np.nan
    if x1.size>1 and x0.size>1:
        m1, m0 = x1.mean(), x0.mean()
        s1, s0 = x1.std(ddof=1), x0.std(ddof=1)
        n1, n0 = x1.size, x0.size
        denom = (n1+n0-2)
        s_pooled = math.sqrt(((n1-1)*s1**2 + (n0-1)*s0**2)/denom) if denom>0 else np.nan
        if s_pooled and not np.isnan(s_pooled) and s_pooled>0:
            d = (m1-m0)/s_pooled
            J = 1 - (3/(4*(n1+n0)-9)) if (n1+n0)>3 else 1.0
            g = J*d

    dur_row = {
        "Variable": "Dauer (Stunden)",
        gmap[1]: f"{r(x1.median())} [{r(x1.quantile(.25))}; {r(x1.quantile(.75))}] (n={x1.size})",
        gmap[0]: f"{r(x0.median())} [{r(x0.quantile(.25))}; {r(x0.quantile(.75))}] (n={x0.size})",
        "Hedges g": r(g,2),
        "p (Mann-Whitney)": r(p_mw,3) if p_mw is not None else "",
        "p (Welch)": r(p_welch,3) if p_welch is not None else "",
        "Hinweis": "Median[IQR]; nichtparametrisch (MW); Welch als Zusatz"
    }
    rows.append(dur_row)

    # ---- Kategoriale Blöcke ----
    def cat_block(col, label):
        out = []
        ser = surv[col]
        if ser.isna().all(): 
            return out
        ctab = pd.crosstab(ser, surv["Gruppe"])
        # Binär?
        if ctab.shape[0]==2:
            OR, ci, RD, p, test = binary_effects(ctab)
            for cat, row in ctab.iterrows():
                n1, n0 = int(row.get(gmap[1],0)), int(row.get(gmap[0],0))
                tot1 = int(ctab[gmap[1]].sum()); tot0 = int(ctab[gmap[0]].sum())
                out.append({
                    "Variable": f"{label}: {cat}",
                    gmap[1]: f"{n1} ({r(100*n1/tot1,1)}%)" if tot1 else f"{n1}",
                    gmap[0]: f"{n0} ({r(100*n0/tot0,1)}%)" if tot0 else f"{n0}",
                    "Effekt": f"OR {r(OR,2)} [{r(ci[0],2)}; {r(ci[1],2)}]",
                    "p-Wert": r(p,3) if p is not None else "",
                    "Test": test
                })
            out.append({
                "Variable": f"{label}: Risikodifferenz (G1−G0, erste Kategorie)",
                gmap[1]: f"n={int(ctab[gmap[1]].sum())}",
                gmap[0]: f"n={int(ctab[gmap[0]].sum())}",
                "Effekt": f"RD {r(RD,3)}",
                "p-Wert": r(p,3) if p is not None else "",
                "Test": test
            })
        else:
            V, p, test = kgt2_effect(ctab)
            for cat, row in ctab.iterrows():
                n1, n0 = int(row.get(gmap[1],0)), int(row.get(gmap[0],0))
                tot1 = int(ctab[gmap[1]].sum()); tot0 = int(ctab[gmap[0]].sum())
                out.append({
                    "Variable": f"{label}: {cat}",
                    gmap[1]: f"{n1} ({r(100*n1/tot1,1)}%)" if tot1 else f"{n1}",
                    gmap[0]: f"{n0} ({r(100*n0/tot0,1)}%)" if tot0 else f"{n0}",
                    "Effekt": f"Cramérs V {r(V,2)}",
                    "p-Wert": r(p,3) if p is not None else "",
                    "Test": test
                })
        return out

    rows += cat_block("Sex_norm", "Geschlecht")
    if "is_reop" in surv.columns:
        rows += cat_block("is_reop", "Re-OP (ja=1, nein=0)")
    if "duration_tertile" in surv.columns:
        rows += cat_block("duration_tertile", "OP-Dauer (Tertile)")

    table = pd.DataFrame(rows)
    return table

def maybe_write_docx(table: pd.DataFrame, path: str):
    try:
        from docx import Document
        doc = Document()
        doc.add_heading('Table 1 (OP-Level): Basischarakteristika mit Tests', level=1)
        doc.add_paragraph('Gruppen: AKI-Index-OP vs. Keine AKI-Index-OP. Prozentwerte sind innerhalb der jeweiligen Gruppe berechnet.')
        doc.add_paragraph('Dauer: Median [IQR]; Mann-Whitney p; Welch p als Zusatz.')
        doc.add_paragraph('Kategorial: Fisher/Chi²; Effekt: OR (95%-KI) bzw. Cramérs V; Risikodifferenz zusätzlich.')
        t = doc.add_table(rows=1, cols=len(table.columns))
        for j, c in enumerate(table.columns):
            t.rows[0].cells[j].text = str(c)
        for _, row in table.iterrows():
            cells = t.add_row().cells
            for j, c in enumerate(table.columns):
                cells[j].text = "" if pd.isna(row[c]) else str(row[c])
        doc.save(path)
        return True
    except Exception:
        return False

def main():
    os.makedirs(OUT_DIR, exist_ok=True)
    adata = read_h5ad(H5)

    table = build_table1(adata)
    table.to_csv(CSV_OUT, index=False)
    print("CSV gespeichert:", CSV_OUT)

    # in uns ablegen
    # nur einfache Typen + Pfad zur CSV speichern
    adata.uns["table1"] = {
    "groups": ["AKI-Index-OP", "Keine AKI-Index-OP"],
    "notes": {
        "duration_tests": "Mann-Whitney (primär), Welch t (Zusatz)",
        "categorical_tests": "Fisher/Chi²; Effekte: OR (2x2) bzw. Cramérs V (>2 Kat.), RD zusätzlich"
    },
    "table_path": CSV_OUT,
    "n_rows": int(table.shape[0]),
    "n_cols": int(table.shape[1]),
    "columns": table.columns.tolist()
}
# optional zusätzlich eine stringifizierte DF
    adata.uns["table1_df"] = table.astype(str)

    # kurze Kontrolle
    ad = read_h5ad(H5)
    print("uns keys:", list(ad.uns.keys()))
    print("table1 rows:", len(ad.uns.get("table1", {}).get("table", [])))

if __name__ == "__main__":
    main()
