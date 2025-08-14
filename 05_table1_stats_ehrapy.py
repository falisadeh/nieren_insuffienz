import os, math
import numpy as np
import pandas as pd
from anndata import read_h5ad

H5 = "/Users/fa/Library/Mobile Documents/com~apple~CloudDocs/cs-transfer/aki_ops_master_S1_survival.h5ad"
OUT_DIR = "/Users/fa/Library/Mobile Documents/com~apple~CloudDocs/cs-transfer/Diagramme"
CSV_OUT = os.path.join(OUT_DIR, "Table1_OP_level_with_stats.csv")
DOCX_OUT = os.path.join(OUT_DIR, "Table1_OP_level_with_stats.docx")

try:
    from scipy import stats
    SCIPY = True
except Exception:
    SCIPY = False

def r(x, nd=2):
    if x is None or (isinstance(x, float) and (np.isnan(x) or np.isinf(x))):
        return ""
    return f"{x:.{nd}f}"

def binary_effects(ctab_2x2: pd.DataFrame):
    a,b = int(ctab_2x2.iloc[0,0]), int(ctab_2x2.iloc[0,1])
    c,d = int(ctab_2x2.iloc[1,0]), int(ctab_2x2.iloc[1,1])
    aa,bb,cc,dd = a or 0.5, b or 0.5, c or 0.5, d or 0.5
    OR = (aa*dd)/(bb*cc)
    se = math.sqrt(1/aa + 1/bb + 1/cc + 1/dd)
    ci_lo, ci_hi = math.exp(math.log(OR)-1.96*se), math.exp(math.log(OR)+1.96*se)
    p1 = a/(a+c) if (a+c)>0 else np.nan
    p0 = b/(b+d) if (b+d)>0 else np.nan
    RD = p1 - p0 if not (np.isnan(p1) or np.isnan(p0)) else np.nan
    if SCIPY:
        try:
            _, p = stats.fisher_exact(ctab_2x2.values)
            test = "Fisher"
        except Exception:
            chi2, p, _, _ = stats.chi2_contingency(ctab_2x2.values, correction=False)
            test = "Chi²"
    else:
        p, test = None, "—"
    return OR, (ci_lo, ci_hi), RD, p, test

def kgt2_effect(ctab: pd.DataFrame):
    if SCIPY:
        chi2, p, _, _ = stats.chi2_contingency(ctab.values, correction=False)
        n = ctab.values.sum()
        k = min(ctab.shape)-1
        V = math.sqrt(chi2/(n*k)) if n>0 and k>0 else np.nan
        return V, p, "Chi²"
    return np.nan, None, "—"

def build_table1(adata) -> pd.DataFrame:
    need = ["event_idx","duration_hours","Sex_norm","is_reop","duration_tertile"]
    for c in need:
        if c not in adata.obs.columns:
            adata.obs[c] = np.nan

    surv = adata.obs[need].copy()
    g1, g0 = "AKI-Index-OP", "Keine AKI-Index-OP"
    surv["Gruppe"] = np.where(surv["event_idx"].astype(int)==1, g1, g0)

    rows = []

    # Dauer (Stunden): MWU + Welch + Hedges g
    x1 = pd.to_numeric(surv.loc[surv["Gruppe"]==g1, "duration_hours"], errors="coerce").dropna()
    x0 = pd.to_numeric(surv.loc[surv["Gruppe"]==g0, "duration_hours"], errors="coerce").dropna()
    if SCIPY and x1.size>1 and x0.size>1:
        U, p_mw = stats.mannwhitneyu(x1, x0, alternative="two-sided")
        t, p_welch = stats.ttest_ind(x1, x0, equal_var=False)
    else:
        p_mw, p_welch = None, None

    g = np.nan
    if x1.size>1 and x0.size>1:
        m1, m0 = x1.mean(), x0.mean()
        s1, s0 = x1.std(ddof=1), x0.std(ddof=1)
        n1, n0 = x1.size, x0.size
        denom = (n1+n0-2)
        s_pooled = math.sqrt(((n1-1)*s1**2 + (n0-1)*s0**2)/denom) if denom>0 else np.nan
        if s_pooled and not np.isnan(s_pooled) and s_pooled>0:
            d = (m1-m0)/s_pooled
            J = 1 - (3/(4*(n1+n0)-9)) if (n1+n0)>3 else 1.0
            g = J*d

    rows.append({
        "Variable": "Dauer (Stunden)",
        g1: f"{r(x1.median())} [{r(x1.quantile(.25))}; {r(x1.quantile(.75))}] (n={x1.size})",
        g0: f"{r(x0.median())} [{r(x0.quantile(.25))}; {r(x0.quantile(.75))}] (n={x0.size})",
        "Hedges g": r(g,2),
        "p (Mann-Whitney)": r(p_mw,3) if p_mw is not None else "",
        "p (Welch)": r(p_welch,3) if p_welch is not None else "",
        "Hinweis": "Median[IQR]; MW primär; Welch als Zusatz"
    })

    def cat_block(col, label):
        out = []
        ser = surv[col]
        if ser.isna().all(): 
            return out
        ctab = pd.crosstab(ser, surv["Gruppe"])
        if ctab.shape[0]==2:
            OR, ci, RD, p, test = binary_effects(ctab)
            for cat, row in ctab.iterrows():
                n1, n0 = int(row.get(g1,0)), int(row.get(g0,0))
                tot1 = int(ctab[g1].sum()); tot0 = int(ctab[g0].sum())
                out.append({
                    "Variable": f"{label}: {cat}",
                    g1: f"{n1} ({r(100*n1/tot1,1)}%)" if tot1 else f"{n1}",
                    g0: f"{n0} ({r(100*n0/tot0,1)}%)" if tot0 else f"{n0}",
                    "Effekt": f"OR {r(OR,2)} [{r(ci[0],2)}; {r(ci[1],2)}]",
                    "p-Wert": r(p,3) if p is not None else "",
                    "Test": test
                })
            out.append({
                "Variable": f"{label}: Risikodifferenz (G1−G0, erste Kategorie)",
                g1: f"n={int(ctab[g1].sum())}",
                g0: f"n={int(ctab[g0].sum())}",
                "Effekt": f"RD {r(RD,3)}",
                "p-Wert": r(p,3) if p is not None else "",
                "Test": test
            })
        else:
            V, p, test = kgt2_effect(ctab)
            for cat, row in ctab.iterrows():
                n1, n0 = int(row.get(g1,0)), int(row.get(g0,0))
                tot1 = int(ctab[g1].sum()); tot0 = int(ctab[g0].sum())
                out.append({
                    "Variable": f"{label}: {cat}",
                    g1: f"{n1} ({r(100*n1/tot1,1)}%)" if tot1 else f"{n1}",
                    g0: f"{n0} ({r(100*n0/tot0,1)}%)" if tot0 else f"{n0}",
                    "Effekt": f"Cramérs V {r(V,2)}",
                    "p-Wert": r(p,3) if p is not None else "",
                    "Test": test
                })
        return out

    rows += cat_block("Sex_norm", "Geschlecht")
    rows += cat_block("is_reop", "Re-OP (ja=1, nein=0)")
    rows += cat_block("duration_tertile", "OP-Dauer (Tertile)")

    return pd.DataFrame(rows)

def maybe_write_docx(table: pd.DataFrame, path: str):
    try:
        from docx import Document
    except Exception:
        return False
    doc = Document()
    doc.add_heading('Table 1 (OP-Level): Basischarakteristika mit Tests', level=1)
    doc.add_paragraph('Gruppen: AKI-Index-OP vs. Keine AKI-Index-OP. Prozentwerte sind gruppenintern.')
    doc.add_paragraph('Dauer: Median [IQR]; Mann-Whitney p; Welch p als Zusatz.')
    doc.add_paragraph('Kategorial: Fisher/Chi²; Effekte: OR (95%-KI) bzw. Cramérs V; RD zusätzlich.')
    t = doc.add_table(rows=1, cols=len(table.columns))
    for j, c in enumerate(table.columns):
        t.rows[0].cells[j].text = str(c)
    for _, row in table.iterrows():
        cells = t.add_row().cells
        for j, c in enumerate(table.columns):
            cells[j].text = "" if pd.isna(row[c]) else str(row[c])
    doc.save(path)
    return True



# ===== Compact-Export für Table 1 (CSV + DOCX) =====
import re
import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

def _fmt_p(x):
    try:
        x = float(x)
        return "<0,001" if x < 0.001 else f"{x:.3f}".replace(".", ",")
    except Exception:
        s = str(x)
        if s.strip() in ("0", "0.0", "0.00"):  # exporter setzte teils 0
            return "<0,001"
        return s

def _std_cols(df):
    """Spalten des vorhandenen table1-DF robust erkennen."""
    colmap = {
        "Variable": ["Variable"],
        "grp1": ["AKI-Index-OP", "AKI_Index_OP", "Gruppe1", "Group1"],
        "grp0": ["Keine AKI-Index-OP", "Keine_AKI_Index_OP", "Gruppe0", "Group0"],
        "Test": ["Test"],
        "Effekt": ["Effekt", "Effect"],
        "CI": ["CI", "95%-KI", "95% KI", "KI"],
        "p_cat": ["p-Wert", "p_value", "p-Value"],                 # für kategorial
        "p_mw": ["p (Mann-Whitney)", "p(Mann-Whitney)"],           # für Dauer (stetig)
        "p_welch": ["p (Welch)", "p(Welch)"],                      # Zusatz
        "Hinweis": ["Hinweis"]
    }
    out = {}
    for std, candidates in colmap.items():
        for c in candidates:
            if c in df.columns:
                out[std] = c
                break
    # Pflichtfelder absichern
    for k in ["Variable", "grp1", "grp0", "Test", "Effekt"]:
        out.setdefault(k, k if k in df.columns else df.columns[min(0, len(df.columns)-1)])
    return out

def _merge_binary_rows(df, cols, prefix, label_pair):
    """Führt z. B. 'Geschlecht: f' + 'Geschlecht: m' zu 1 Zeile zusammen."""
    mask = df[cols["Variable"]].astype(str).str.startswith(prefix)
    sub = df.loc[mask]
    if len(sub) != 2:
        return None
    g1 = f'{sub.iloc[0][cols["grp1"]]} / {sub.iloc[1][cols["grp1"]]}'
    g0 = f'{sub.iloc[0][cols["grp0"]]} / {sub.iloc[1][cols["grp0"]]}'
    eff = str(sub.iloc[0].get(cols["Effekt"], ""))
    ci  = str(sub.iloc[0].get(cols.get("CI",""), ""))
    p   = sub.iloc[0].get(cols.get("p_cat", ""), "")
    row = {
        "Variable": f"{prefix} ({label_pair})",
        "AKI-Index-OP": g1,
        "Keine AKI-Index-OP": g0,
        "Test": sub.iloc[0].get(cols["Test"], ""),
        "Effekt (95%-KI)": (eff + (f" {ci}" if ci and ci != "nan" else "")).strip(),
        "p-Wert": _fmt_p(p),
    }
    return row
#---------------------
import re
import numpy as np

def _num(x):
    try:
        return float(str(x).replace(",", "."))
    except Exception:
        return np.nan

def _median_from_cell(cell):
    """
    Holt den ersten Zahlenwert aus Zellen wie '4.5 [3.3; 5.8]' oder 'Median [IQR]'.
    Fällt zurück auf NaN, wenn nichts passt.
    """
    s = str(cell)
    m = re.search(r"[-+]?\d+(?:[.,]\d+)?", s)
    return _num(m.group(0)) if m else np.nan

def _merge_any_binary(df, cols, prefixes, label_pair):
    """
    Findet zwei Zeilen, deren 'Variable' mit einem der Prefixe beginnt (z. B. 'Geschlecht', 'Sex'),
    und führt sie zu einer zusammen. RD-Zeilen werden ignoriert.
    """
    varcol = cols["Variable"]
    series = df[varcol].astype(str)
    mask = False
    for p in prefixes:
        mask = (series.str.startswith(p)) | mask
    sub = df.loc[mask & ~series.str.contains("Risikodifferenz", na=False)]
    if len(sub) != 2:
        return None

    grp1 = f'{sub.iloc[0][cols["grp1"]]} / {sub.iloc[1][cols["grp1"]]}'
    grp0 = f'{sub.iloc[0][cols["grp0"]]} / {sub.iloc[1][cols["grp0"]]}'

    eff = str(sub.iloc[0].get(cols.get("Effekt", ""), ""))
    ci  = str(sub.iloc[0].get(cols.get("CI", ""), ""))
    p   = sub.iloc[0].get(cols.get("p_cat", ""), "")

    return {
        "Variable": f"{prefixes[0]} ({label_pair})",
        "AKI-Index-OP": grp1,
        "Keine AKI-Index-OP": grp0,
        "Test": sub.iloc[0].get(cols.get("Test", ""), ""),
        "Effekt (95%-KI)": (eff + (f" {ci}" if ci and ci != "nan" else "")).strip(),
        "p-Wert": _fmt_p(p),
    }



#-------------------

def make_table1_compact(df_raw: pd.DataFrame) -> pd.DataFrame:
    """Erzeugt eine kompakte Table 1 (eine Zeile pro Variable, ohne RD-Zeilen)."""
    cols = _std_cols(df_raw)
    # Variable-Namen normalisieren: "Geschlecht: f" -> "Geschlecht", "Re-OP (ja=1, nein=0): 0.0" -> "Re-OP"
    df_norm = df_raw.copy()
    varcol = cols["Variable"]
    df_norm["_var_norm"] = (
        df_norm[varcol].astype(str)
        .str.replace(r":.*$", "", regex=True)   # alles nach ":" abschneiden
        .str.replace(r"\s*\(.*\)$", "", regex=True)  # Klammerzusatz entfernen
    )

    def _merge_by_norm(prefixes, label_pair):
        m = df_norm["_var_norm"].isin(prefixes)
        sub = df_norm.loc[m & ~df_norm[varcol].astype(str).str.contains("Risikodifferenz", na=False)]
        if len(sub) != 2:
            return None
        g1 = f'{sub.iloc[0][cols["grp1"]]} / {sub.iloc[1][cols["grp1"]]}'
        g0 = f'{sub.iloc[0][cols["grp0"]]} / {sub.iloc[1][cols["grp0"]]}'
        eff = str(sub.iloc[0].get(cols.get("Effekt",""), ""))
        ci  = str(sub.iloc[0].get(cols.get("CI",""), ""))
        p   = sub.iloc[0].get(cols.get("p_cat",""), "")
        return {
            "Variable": f"{prefixes[0]} ({label_pair})",
            "AKI-Index-OP": g1,
            "Keine AKI-Index-OP": g0,
            "Test": sub.iloc[0].get(cols.get("Test",""), ""),
            "Effekt (95%-KI)": (eff + (f" {ci}" if ci and ci != "nan" else "")).strip(),
            "p-Wert": _fmt_p(p),
        }
    rows = []
    # 1) Dauer (stetig) – p aus MW (primär) oder ersatzweise Welch
    for _, r in df_raw.iterrows():
        v = str(r[cols["Variable"]])
        #if v.startswith("Dauer"):
        if re.match(r'^(Dauer|OP[- ]?Dauer)', v):
            p_val = r.get(cols.get("p_mw",""), "")
            if (isinstance(p_val, float) and np.isnan(p_val)) or p_val in ("", None):
                p_val = r.get(cols.get("p_welch",""), "")
        # Median berechnen
        med1 = _median_from_cell(r.get(cols["grp1"], ""))
        med0 = _median_from_cell(r.get(cols["grp0"], ""))
        dmed = med1 - med0 if np.isfinite(med1) and np.isfinite(med0) else np.nan
        eff_ci = f"Median {dmed:.2f}" if np.isfinite(dmed) else eff_ci
        eff = str(r.get(cols.get("Effekt",""), ""))
        ci  = str(r.get(cols.get("CI",""), ""))
        eff_ci = (eff + (f" {ci}" if ci and ci != "nan" else "")).strip()

        if np.isfinite(dmed):
            eff_ci = f"Median {dmed:.2f}"

        rows.append({
            "Variable": v,
            "AKI-Index-OP": str(r.get(cols["grp1"], "")),
            "Keine AKI-Index-OP": str(r.get(cols["grp0"], "")),
            "Test": "MW (Welch Zusatz)" if cols.get("p_welch") in df_raw.columns else "MW",
            "Effekt (95%-KI)": eff_ci,
            "p-Wert": _fmt_p(p_val),
        })


    # 2) Geschlecht zusammenführen
    #m = _merge_binary_rows(df_raw, cols, "Geschlecht", "w/m")
    m = _merge_by_norm(["Geschlecht", "Sex"], "w/m")
    if m: rows.append(m)

    # 3) Re-OP zusammenführen (ja/nein)
    #m = _merge_binary_rows(df_raw, cols, "Re-OP", "ja/nein")
    m = _merge_by_norm(["Re-OP", "ReOP", "Reoperation", "Re-Operation"], "ja/nein")
    if m: rows.append(m)

    # 4) Restliche Variablen: ohne RD-Zeilen und ohne schon behandelte Prefixe
    handled = ("Dauer", "Geschlecht", "Re-OP")
    for _, r in df_raw.iterrows():
        v = str(r[cols["Variable"]])
        if v.startswith(handled) or "Risikodifferenz" in v:
            continue
        rows.append({
            "Variable": v,
            "AKI-Index-OP": str(r.get(cols["grp1"], "")),
            "Keine AKI-Index-OP": str(r.get(cols["grp0"], "")),
            "Test": str(r.get(cols["Test"], "")),
            "Effekt (95%-KI)": str(r.get(cols["Effekt"], "")) + (f' {r.get(cols.get("CI",""), "")}' if cols.get("CI") in df_raw.columns else ""),
            "p-Wert": _fmt_p(r.get(cols.get("p_cat",""), "")),
        })

    compact = pd.DataFrame(rows, columns=[
        "Variable", "AKI-Index-OP", "Keine AKI-Index-OP", "Test", "Effekt (95%-KI)", "p-Wert"
    ])
    return compact

def write_table1_compact_docx(df_compact: pd.DataFrame, path_docx: str):
    """Schreibt die kompakte Table 1 hübsch nach DOCX."""
    doc = Document()
    doc.add_heading("Table 1 (OP-Level): Basiskohorte mit Tests", level=2)

    cols = ["Variable","AKI-Index-OP","Keine AKI-Index-OP","Test","Effekt (95%-KI)","p-Wert"]
    tbl = doc.add_table(rows=1, cols=len(cols))
    tbl.style = "Table Grid"
    for i, c in enumerate(cols):
        tbl.cell(0, i).text = c

    for _, r in df_compact.iterrows():
        row = tbl.add_row().cells
        row[0].text = str(r["Variable"])
        row[1].text = str(r["AKI-Index-OP"])
        row[2].text = str(r["Keine AKI-Index-OP"])
        row[3].text = str(r["Test"])
        row[4].text = str(r["Effekt (95%-KI)"])
        row[5].text = str(r["p-Wert"])

    # Spaltenbreiten & Ausrichtung
    widths_cm = [6, 4, 4, 3, 5, 2.2]
    for i, w in enumerate(widths_cm):
        for row in tbl.rows:
            row.cells[i].width = Cm(w)
            row.cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.save(path_docx)

    #---------------------
def main():
    os.makedirs(OUT_DIR, exist_ok=True)

    # 1) Daten laden
    adata = read_h5ad(H5)

    # 2) Table 1 (raw) bauen + speichern
    table_raw = build_table1(adata)
    table_raw.to_csv(CSV_OUT, index=False)
    print("CSV (raw) gespeichert:", CSV_OUT)

    # 3) Kompakte Table 1 aus der Rohfassung erzeugen
    compact = make_table1_compact(table_raw)

    # --- Helper lokal: 2 Zeilen aus RAW (Geschlecht/Re-OP) mergen ---
    def _merge_from_raw(df_raw, prefix, label_pair):
        vr = df_raw["Variable"].astype(str)
        sub = df_raw[ vr.str.startswith(prefix) & ~vr.str.contains("Risikodifferenz", na=False) ]
        if len(sub) != 2:
            return None
        return {
            "Variable":            f"{prefix} ({label_pair})",
            "AKI-Index-OP":        f"{sub.iloc[0]['AKI-Index-OP']} / {sub.iloc[1]['AKI-Index-OP']}",
            "Keine AKI-Index-OP":  f"{sub.iloc[0]['Keine AKI-Index-OP']} / {sub.iloc[1]['Keine AKI-Index-OP']}",
            "Test":                str(sub.iloc[0].get("Test","")),
            "Effekt (95%-KI)":     str(sub.iloc[0].get("Effekt (95%-KI)", "")),
            "p-Wert":              _fmt_p(sub.iloc[0].get("p-Wert","")),
        }

    # 3a) Fehlende Zeilen ergänzen (Geschlecht & Re-OP)
    rows_add = []
    for pref, label in [("Geschlecht", "w/m"), ("Re-OP", "ja/nein")]:
        row = _merge_from_raw(table_raw, pref, label)
        if row:
            rows_add.append(row)
    if rows_add:
        compact = pd.concat([compact, pd.DataFrame(rows_add)], ignore_index=True)

    # 3b) ΔMedian für "Dauer" setzen, falls noch NaN
    def _med(cell):
        m = re.search(r"[-+]?\d+(?:[.,]\d+)?", str(cell))
        return float(m.group(0).replace(",", ".")) if m else np.nan

    mask_dauer = compact["Variable"].astype(str).str.match(r"^(Dauer|OP[- ]?Dauer)")
    if mask_dauer.any():
        eff_col = "Effekt (95%-KI)"
        if compact.loc[mask_dauer, eff_col].isna().any():
            med1 = _med(compact.loc[mask_dauer, "AKI-Index-OP"].iloc[0])
            med0 = _med(compact.loc[mask_dauer, "Keine AKI-Index-OP"].iloc[0])
            if np.isfinite(med1) and np.isfinite(med0):
                compact.loc[mask_dauer, eff_col] = f"ΔMedian {med1 - med0:.2f}"

    # 4) Aufräumen (keine RD-/Roh-Zeilen; Tertile-Doppel entfernen)
    v = compact["Variable"].astype(str)
    drop_mask = (
        v.str.startswith("Geschlecht:") |
        v.str.contains(r"^Re-OP.*:", regex=True) |
        v.str.contains("Risikodifferenz", na=False)
    )
    compact = compact.loc[~drop_mask].copy()

    mask_tert = compact["Variable"].astype(str).str.startswith("OP-Dauer (Tertile)")
    mask_mw   = compact["Test"].astype(str).str.contains("MW", na=False)
    compact   = compact.loc[~(mask_tert & mask_mw)].copy()

    compact["Effekt (95%-KI)"] = (
        compact["Effekt (95%-KI)"].astype(str).str.replace("Median ", "ΔMedian ", n=1)
    )

    # 5) Reihenfolge robust festlegen (dup-sicher, ohne reindex-Fehler)
    order = ["Dauer (Stunden)", "Geschlecht (w/m)", "Re-OP (ja/nein)"]
    dedup = compact.drop_duplicates(subset=["Variable"], keep="first").copy()

    top_df  = dedup[dedup["Variable"].isin(order)].copy()
    # gewünschte Reihenfolge erzwingen mittels Kategorie
    #cat = pd.Categorical(top_df["Variable"], categories=order, ordered=True)
    #top_df = top_df.assign(_ord=cat).sort_values("_ord").drop(columns="_ord")
    top_df["__ord"] = pd.Categorical(top_df["Variable"], categories=order, ordered=True)
    top_df = top_df.sort_values("__ord").drop(columns="__ord")

    rest_df = dedup[~dedup["Variable"].isin(order)].copy()
    compact = pd.concat([top_df, rest_df], ignore_index=True)

    # 6) Kompakte CSV + DOCX speichern
    CSV_COMPACT  = os.path.join(OUT_DIR, "Table1_OP_level_compact.csv")
    DOCX_COMPACT = os.path.join(OUT_DIR, "Table1_OP_level_compact.docx")
    compact.to_csv(CSV_COMPACT, index=False)
    write_table1_compact_docx(compact, DOCX_COMPACT)
    print("CSV (compact) gespeichert:", CSV_COMPACT)
    print("DOCX (compact) gespeichert:", DOCX_COMPACT)

    # 7) H5AD ohne uns['table1'] schreiben + Kontrolle
    adata.write_h5ad(H5)
    print(f"uns keys: {list(read_h5ad(H5).uns.keys())}")




if __name__ == "__main__":
    main()
